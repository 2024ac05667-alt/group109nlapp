import sys
import subprocess
import os

try:
    from docx import Document
except ImportError:
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
    from docx import Document

report_path = os.path.join(os.path.dirname(__file__), '..', 'Group109_Report.docx')
report_path = os.path.abspath(report_path)

doc = Document()
doc.add_heading('Group 109 Report: KG-Augmented Generative QA', level=1)

doc.add_heading('1. Title & Team Details', level=2)
doc.add_paragraph('Project: Knowledge Graph-Augmented Generative QA')
doc.add_paragraph('Group: 109')
doc.add_paragraph('Team members:')
doc.add_paragraph('- Name 1 (Roll No.)')
doc.add_paragraph('- Name 2 (Roll No.)')
doc.add_paragraph('- Name 3 (Roll No.)')

doc.add_heading('2. Objective', level=2)
doc.add_paragraph(
    'This project builds a lightweight knowledge graph-augmented question answering system using Neo4j and a local open-source language model. ' 
    'It ingests text passages into a Neo4j knowledge graph, translates natural language questions into Cypher queries, retrieves supporting facts, and generates answers using a local LLM.'
)

doc.add_heading('3. Project Components', level=2)
doc.add_paragraph('The repository contains the following components:')
for item in [
    'backend/ - FastAPI application',
    'kg/ - Neo4j client and ingestion logic',
    'llm/ - answer generation module',
    'scripts/ - data preparation and ingestion scripts',
    'data/org_passages.csv - generated corpus',
    'README.md - setup and usage instructions',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4. Setup Instructions', level=2)
doc.add_paragraph('Environment: Python 3.9+ and Neo4j Desktop (or another running Neo4j instance).')
doc.add_paragraph('Steps:')
doc.add_paragraph('1. Create and activate a virtual environment.', style='List Number')
doc.add_paragraph('2. Install required packages: python -m pip install -r requirements.txt', style='List Number')
doc.add_paragraph('3. (Optional) Install Streamlit for UI: python -m pip install streamlit', style='List Number')
doc.add_paragraph('4. Start Neo4j and set environment variables: NEO4J_URI, NEO4J_USER, NEO4J_PASSWORD, NEO4J_DATABASE', style='List Number')

doc.add_heading('5. LLM Details', level=2)
doc.add_paragraph('The project uses the open-source GPT-2 model via Hugging Face transformers. GPT-2 is free to download and run locally. The model files are not included in requirements and are downloaded automatically on first use.')

doc.add_heading('6. How to Run', level=2)
doc.add_paragraph('The main commands are:')
doc.add_paragraph('1. python prepare_data.py', style='List Number')
doc.add_paragraph('2. python build_index.py', style='List Number')
doc.add_paragraph('3. python -m uvicorn backend.main:app --reload --port 8000', style='List Number')
doc.add_paragraph('4. Optional: streamlit run app.py', style='List Number')
doc.add_paragraph('5. Query the API using curl or browser: http://127.0.0.1:8000/query?question=...', style='List Number')

doc.add_heading('7. Dataset Description', level=2)
doc.add_paragraph('The dataset file data/org_passages.csv contains synthetic corporate passages used to populate the knowledge graph. It includes organizational facts and relationships for the QA system.')

doc.add_heading('8. Results', level=2)
doc.add_paragraph('Sample queries and outputs should be captured during execution. Example: asking "Who is the CEO of the company?" returns a generated answer along with retrieved supporting facts and the translated Cypher query.')

doc.add_heading('9. Screenshots / Evidence', level=2)
doc.add_paragraph('Include screenshots showing:')
doc.add_paragraph('- Neo4j Desktop with the corporg database running', style='List Bullet')
doc.add_paragraph('- build_index.py completion', style='List Bullet')
doc.add_paragraph('- backend running on uvicorn', style='List Bullet')
doc.add_paragraph('- sample query response from the API or Streamlit UI', style='List Bullet')

doc.add_heading('10. Limitations / Notes', level=2)
doc.add_paragraph('The NL-to-Cypher translator is rule-based and works best for simple question patterns. The default GPT-2 model is lightweight and may produce basic answers. spaCy and Streamlit are optional and can be skipped if the environment does not support them.')

doc.add_heading('11. Files Included in ZIP', level=2)
doc.add_paragraph('The ZIP file should include:')
doc.add_paragraph('- README.md', style='List Bullet')
doc.add_paragraph('- requirements.txt', style='List Bullet')
doc.add_paragraph('- app.py, build_index.py, evaluate.py, prepare_data.py, nl2cypher.py', style='List Bullet')
doc.add_paragraph('- backend/, kg/, llm/, scripts/', style='List Bullet')
doc.add_paragraph('- data/org_passages.csv', style='List Bullet')
doc.add_paragraph('- eval/sample_queries.txt', style='List Bullet')

doc.save(report_path)
print(report_path)
